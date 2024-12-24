import pdfplumber
from docx import Document


def ReadDocx(file_path):
    document = Document(file_path)

    for para in document.paragraphs:
        if "Навыки" in para.text:

            print(f"Найден текст: '{para.text}'")
    jobName = ""
    # Ищем в таблицах
    for table in document.tables:
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                if "Навыки" in cell.text:
                    if cell.text != row.cells[1].text:
                        skills = row.cells[1].text

                if "должность" in cell.text:
                    if row_index + 1 < len(table.rows):
                        next_row = table.rows[row_index + 1]
                        if next_row.cells:
                            jobName = next_row.cells[0].text.splitlines()[0]

    return jobName


def ReadPDF(file_path):
    jobName = ""

    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            # Извлекаем текст со страницы
            text = page.extract_text()
            if text:
                # Разделяем текст на строки
                lines = text.splitlines()

                # Ищем "Навыки" в тексте
                for line in lines:
                    if "Навыки" in line:
                        print(f"Найден текст: '{line}'")

                # Ищем "должность"
                for i, line in enumerate(lines):
                    if "должность" in line:
                        # Проверяем, есть ли следующая строка
                        if i + 1 < len(lines):
                            jobName = lines[i + 1].splitlines()[0]
                            break  # Выходим из цикла, если нашли должность

    return jobName
